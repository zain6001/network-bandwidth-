"""
Report Generator Module
Generates network traffic reports in Markdown and Word formats
"""

from datetime import datetime
import os


class ReportGenerator:
    def __init__(self, database_logger):
        self.db = database_logger
    
    def generate_markdown_report(self, start_date, end_date, output_file):
        """Generate a detailed markdown report"""
        try:
            # Get data from database
            overall_stats, process_stats = self.db.get_detailed_report_data(start_date, end_date)
            alerts = self.db.get_alerts_by_date_range(start_date, end_date)
            sessions = self.db.get_sessions_in_range(start_date, end_date)
            
            # Generate report content
            report = []
            report.append("# Network Traffic Report\n")
            report.append(f"**Report Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            report.append(f"**Period:** {start_date.strftime('%Y-%m-%d %H:%M')} to {end_date.strftime('%Y-%m-%d %H:%M')}\n")
            report.append("\n---\n\n")
            
            # Overall Statistics
            report.append("## Overall Statistics\n\n")
            if overall_stats:
                unique_procs, total_up, total_down, avg_up, avg_down, max_up, max_down = overall_stats
                
                report.append(f"- **Unique Processes:** {unique_procs or 0}\n")
                report.append(f"- **Total Upload:** {(total_up or 0) / (1024*1024):.2f} MB\n")
                report.append(f"- **Total Download:** {(total_down or 0) / (1024*1024):.2f} MB\n")
                report.append(f"- **Total Data:** {((total_up or 0) + (total_down or 0)) / (1024*1024):.2f} MB\n")
                report.append(f"- **Average Upload Rate:** {(avg_up or 0) / 1024:.2f} KB/s\n")
                report.append(f"- **Average Download Rate:** {(avg_down or 0) / 1024:.2f} KB/s\n")
                report.append(f"- **Peak Upload Rate:** {(max_up or 0) / 1024:.2f} KB/s\n")
                report.append(f"- **Peak Download Rate:** {(max_down or 0) / 1024:.2f} KB/s\n")
            else:
                report.append("*No data available for this period*\n")
            
            report.append("\n")
            
            # Monitoring Sessions
            if sessions:
                report.append("## Monitoring Sessions\n\n")
                report.append(f"Total sessions: {len(sessions)}\n\n")
                report.append("| Session ID | Start Time | End Time | Upload | Download |\n")
                report.append("|------------|------------|----------|---------|----------|\n")
                for session in sessions[:10]:  # Show latest 10
                    sid, start, end, up, down = session
                    start_time = datetime.fromisoformat(start).strftime('%Y-%m-%d %H:%M')
                    end_time = datetime.fromisoformat(end).strftime('%H:%M') if end else 'Running'
                    report.append(f"| {sid} | {start_time} | {end_time} | {(up or 0)/(1024*1024):.2f} MB | {(down or 0)/(1024*1024):.2f} MB |\n")
                report.append("\n")
            
            # Per-Process Statistics
            report.append("## Per-Process Traffic\n\n")
            if process_stats:
                report.append("| PID | Process Name | Upload | Download | Total | Avg Up Rate | Avg Down Rate | Max Up Rate | Max Down Rate | Protocols |\n")
                report.append("|-----|--------------|--------|----------|-------|-------------|---------------|-------------|---------------|----------|\n")
                
                for proc in process_stats:
                    pid, name, up, down, avg_up, avg_down, max_up, max_down, records, protocols = proc
                    total = up + down
                    report.append(f"| {pid} | {name} | {up/(1024*1024):.2f} MB | {down/(1024*1024):.2f} MB | {total/(1024*1024):.2f} MB | ")
                    report.append(f"{avg_up/1024:.1f} KB/s | {avg_down/1024:.1f} KB/s | {max_up/1024:.1f} KB/s | {max_down/1024:.1f} KB/s | {protocols or 'N/A'} |\n")
            else:
                report.append("*No process data available*\n")
            
            report.append("\n")
            
            # Bandwidth Spikes / Alerts
            report.append("## Bandwidth Alerts & Spikes\n\n")
            if alerts:
                report.append(f"Total alerts: {len(alerts)}\n\n")
                report.append("| Timestamp | PID | Process | Alert Type | Bandwidth | Threshold |\n")
                report.append("|-----------|-----|---------|------------|-----------|----------|\n")
                
                for alert in alerts:
                    timestamp, pid, name, alert_type, bandwidth, threshold = alert
                    time_str = datetime.fromisoformat(timestamp).strftime('%Y-%m-%d %H:%M:%S')
                    report.append(f"| {time_str} | {pid} | {name} | {alert_type} | {bandwidth:.1f} KB/s | {threshold:.1f} KB/s |\n")
            else:
                report.append("*No bandwidth alerts during this period*\n")
            
            report.append("\n")
            
            # Top Consumers
            if process_stats:
                report.append("## Top 5 Bandwidth Consumers\n\n")
                top_5 = sorted(process_stats, key=lambda x: x[2] + x[3], reverse=True)[:5]
                
                for i, proc in enumerate(top_5, 1):
                    pid, name, up, down, _, _, _, _, _, _ = proc
                    total = (up + down) / (1024 * 1024)
                    report.append(f"{i}. **{name}** (PID {pid}): {total:.2f} MB total\n")
                    report.append(f"   - Upload: {up/(1024*1024):.2f} MB\n")
                    report.append(f"   - Download: {down/(1024*1024):.2f} MB\n")
                    report.append("\n")
            
            # Write to file
            with open(output_file, 'w') as f:
                f.writelines(report)
            
            return True, output_file
            
        except Exception as e:
            return False, str(e)
    
    def generate_word_report(self, start_date, end_date, output_file):
        """Generate a Word document report using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Get data
            overall_stats, process_stats = self.db.get_detailed_report_data(start_date, end_date)
            alerts = self.db.get_alerts_by_date_range(start_date, end_date)
            sessions = self.db.get_sessions_in_range(start_date, end_date)
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Network Traffic Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Report info
            info = doc.add_paragraph()
            info.add_run('Report Generated: ').bold = True
            info.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            info.add_run('Period: ').bold = True
            info.add_run(f"{start_date.strftime('%Y-%m-%d %H:%M')} to {end_date.strftime('%Y-%m-%d %H:%M')}")
            
            doc.add_paragraph('_' * 50)
            
            # Overall Statistics
            doc.add_heading('Overall Statistics', 1)
            if overall_stats:
                unique_procs, total_up, total_down, avg_up, avg_down, max_up, max_down = overall_stats
                
                stats_table = doc.add_table(rows=8, cols=2)
                stats_table.style = 'Light Grid Accent 1'
                
                stats_data = [
                    ('Unique Processes', f"{unique_procs or 0}"),
                    ('Total Upload', f"{(total_up or 0) / (1024*1024):.2f} MB"),
                    ('Total Download', f"{(total_down or 0) / (1024*1024):.2f} MB"),
                    ('Total Data', f"{((total_up or 0) + (total_down or 0)) / (1024*1024):.2f} MB"),
                    ('Average Upload Rate', f"{(avg_up or 0) / 1024:.2f} KB/s"),
                    ('Average Download Rate', f"{(avg_down or 0) / 1024:.2f} KB/s"),
                    ('Peak Upload Rate', f"{(max_up or 0) / 1024:.2f} KB/s"),
                    ('Peak Download Rate', f"{(max_down or 0) / 1024:.2f} KB/s"),
                ]
                
                for i, (label, value) in enumerate(stats_data):
                    row = stats_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = value
            
            # Monitoring Sessions
            if sessions:
                doc.add_heading('Monitoring Sessions', 1)
                doc.add_paragraph(f'Total sessions: {len(sessions)}')
                
                sess_table = doc.add_table(rows=1, cols=5)
                sess_table.style = 'Light Grid Accent 1'
                hdr_cells = sess_table.rows[0].cells
                hdr_cells[0].text = 'Session ID'
                hdr_cells[1].text = 'Start Time'
                hdr_cells[2].text = 'End Time'
                hdr_cells[3].text = 'Upload'
                hdr_cells[4].text = 'Download'
                
                for session in sessions[:10]:
                    sid, start, end, up, down = session
                    row = sess_table.add_row().cells
                    row[0].text = str(sid)
                    row[1].text = datetime.fromisoformat(start).strftime('%Y-%m-%d %H:%M')
                    row[2].text = datetime.fromisoformat(end).strftime('%H:%M') if end else 'Running'
                    row[3].text = f"{(up or 0)/(1024*1024):.2f} MB"
                    row[4].text = f"{(down or 0)/(1024*1024):.2f} MB"
            
            # Per-Process Statistics
            doc.add_heading('Per-Process Traffic', 1)
            if process_stats:
                proc_table = doc.add_table(rows=1, cols=6)
                proc_table.style = 'Light Grid Accent 1'
                hdr_cells = proc_table.rows[0].cells
                hdr_cells[0].text = 'PID'
                hdr_cells[1].text = 'Process'
                hdr_cells[2].text = 'Upload'
                hdr_cells[3].text = 'Download'
                hdr_cells[4].text = 'Total'
                hdr_cells[5].text = 'Protocols'
                
                for proc in process_stats:
                    pid, name, up, down, _, _, _, _, _, protocols = proc
                    row = proc_table.add_row().cells
                    row[0].text = str(pid)
                    row[1].text = name
                    row[2].text = f"{up/(1024*1024):.2f} MB"
                    row[3].text = f"{down/(1024*1024):.2f} MB"
                    row[4].text = f"{(up+down)/(1024*1024):.2f} MB"
                    row[5].text = protocols or 'N/A'
            
            # Bandwidth Alerts
            doc.add_heading('Bandwidth Alerts & Spikes', 1)
            if alerts:
                doc.add_paragraph(f'Total alerts: {len(alerts)}')
                
                alert_table = doc.add_table(rows=1, cols=6)
                alert_table.style = 'Light Grid Accent 1'
                hdr_cells = alert_table.rows[0].cells
                hdr_cells[0].text = 'Timestamp'
                hdr_cells[1].text = 'PID'
                hdr_cells[2].text = 'Process'
                hdr_cells[3].text = 'Type'
                hdr_cells[4].text = 'Bandwidth'
                hdr_cells[5].text = 'Threshold'
                
                for alert in alerts:
                    timestamp, pid, name, alert_type, bandwidth, threshold = alert
                    row = alert_table.add_row().cells
                    row[0].text = datetime.fromisoformat(timestamp).strftime('%Y-%m-%d %H:%M:%S')
                    row[1].text = str(pid)
                    row[2].text = name
                    row[3].text = alert_type
                    row[4].text = f"{bandwidth:.1f} KB/s"
                    row[5].text = f"{threshold:.1f} KB/s"
            else:
                doc.add_paragraph('No bandwidth alerts during this period')
            
            # Top Consumers
            if process_stats:
                doc.add_heading('Top 5 Bandwidth Consumers', 1)
                top_5 = sorted(process_stats, key=lambda x: x[2] + x[3], reverse=True)[:5]
                
                for i, proc in enumerate(top_5, 1):
                    pid, name, up, down, _, _, _, _, _, _ = proc
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f"{name} ").bold = True
                    p.add_run(f"(PID {pid}): {(up+down)/(1024*1024):.2f} MB total")
            
            # Save document
            doc.save(output_file)
            return True, output_file
            
        except ImportError:
            return False, "python-docx not installed. Install with: pip install python-docx"
        except Exception as e:
            return False, str(e)
